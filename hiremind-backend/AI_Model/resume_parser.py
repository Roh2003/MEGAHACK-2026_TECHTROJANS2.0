from __future__ import annotations

import io
from pathlib import Path

import PyPDF2
import requests
from docx import Document


class ResumeParser:
    """Extract plain text from resume files using local paths or URLs."""

    def __init__(self, base_dir: str | Path | None = None, timeout: int = 20) -> None:
        self.base_dir = Path(base_dir).resolve() if base_dir else Path.cwd()
        self.timeout = timeout

    def parse(self, source: str) -> str:
        path_or_url = (source or "").strip()
        if not path_or_url:
            return ""

        if path_or_url.lower().startswith(("http://", "https://")):
            return self._parse_url(path_or_url)

        return self._parse_local(path_or_url)

    def _parse_url(self, url: str) -> str:
        response = requests.get(url, timeout=self.timeout)
        response.raise_for_status()

        content_type = (response.headers.get("Content-Type") or "").lower()
        ext = self._guess_extension(url, content_type)
        data = response.content

        if ext == "pdf":
            return self._pdf_bytes_to_text(data)
        if ext == "txt":
            return data.decode("utf-8", errors="ignore")
        if ext == "docx":
            return self._docx_bytes_to_text(data)

        return ""

    def _parse_local(self, path_text: str) -> str:
        candidate_path = Path(path_text)
        if not candidate_path.is_absolute():
            candidate_path = (self.base_dir / candidate_path).resolve()

        if not candidate_path.exists() or not candidate_path.is_file():
            return ""

        ext = candidate_path.suffix.lower().replace(".", "")

        if ext == "pdf":
            return self._pdf_bytes_to_text(candidate_path.read_bytes())
        if ext == "txt":
            return candidate_path.read_text(encoding="utf-8", errors="ignore")
        if ext == "docx":
            return self._docx_path_to_text(candidate_path)

        return ""

    @staticmethod
    def _guess_extension(url: str, content_type: str) -> str:
        lower_url = url.lower().split("?")[0]
        if lower_url.endswith(".pdf") or "pdf" in content_type:
            return "pdf"
        if lower_url.endswith(".txt") or "text/plain" in content_type:
            return "txt"
        if lower_url.endswith(".docx") or "wordprocessingml" in content_type:
            return "docx"
        return ""

    @staticmethod
    def _pdf_bytes_to_text(data: bytes) -> str:
        reader = PyPDF2.PdfReader(io.BytesIO(data))
        chunks: list[str] = []
        for page in reader.pages:
            chunks.append(page.extract_text() or "")
        return "\n".join(chunks).strip()

    @staticmethod
    def _docx_bytes_to_text(data: bytes) -> str:
        doc = Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs if p.text).strip()

    @staticmethod
    def _docx_path_to_text(path: Path) -> str:
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs if p.text).strip()
